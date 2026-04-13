"""
Export & Rendering Integration Tests
=====================================
Tests TXT export whitespace cleanup, Word report generation,
source renumbering, and evidence filtering — without Streamlit or API calls.
"""

import re
import io
from docx import Document


# ---------------------------------------------------------------------------
# Helpers (extracted logic from app_render.py, testable without Streamlit)
# ---------------------------------------------------------------------------

WHITESPACE_RE = re.compile(r'[^\S\n]{2,}')


def clean_export_text(text: str) -> str:
    """Same regex used in app_render.py for TXT exports."""
    return WHITESPACE_RE.sub(' ', text)


def build_renumber_map(answer: str) -> dict[int, int]:
    """Same logic as app_render.py: internal chunk IDs → sequential display IDs."""
    cited_raw = sorted(set(int(x) for x in re.findall(r'\[(\d+)\]', answer)))
    return {old: new for new, old in enumerate(cited_raw, start=1)}


def renumber_text(text: str, renumber: dict[int, int]) -> str:
    """Apply renumbering to citation markers in text."""
    if not renumber:
        return text
    return re.sub(
        r'\[(\d+)\]',
        lambda m: f"[{renumber[int(m.group(1))]}]"
        if int(m.group(1)) in renumber else m.group(0),
        text
    )


def filter_cited_refs(refs: list[dict], answer: str) -> list[dict]:
    """Same filtering as render_evidence_list: only show cited sources."""
    displayed = [r for r in refs if f"[{r['id']}]" in answer]
    if not displayed:
        displayed = refs[:3]
    return displayed


def build_all_sources_txt(refs: list[dict]) -> str:
    """Builds the 'Download All Sources' TXT content."""
    header = f"DOCUINSIGHT - ALL SOURCES\nTotal sources: {len(refs)}\n{'='*50}\n\n"
    body = ""
    for r in refs:
        clean = clean_export_text(r.get('text', '[No Text]'))
        body += (
            f"{'='*50}\n"
            f"SOURCE [{r['id']}]\n"
            f"File: {r['file']}\n"
            f"Page: {r['page']}\n"
            f"{'='*50}\n\n"
            f"{clean}\n\n"
        )
    return header + body


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

# Simulates PyPDFLoader output with various whitespace artifacts
PDF_TEXT_WITH_ARTIFACTS = (
    "Received:\t\t21\t\tMarch\t\t2023  |  Revised:\t\t14\t\tJuly\t\t2023\n"
    "AD,\u00a0\u00a0\u00a0\u00a0Alzheimer's\u00a0\u00a0\u00a0\u00a0disease;\u00a0\u00a0\u00a0\u00a0AKI,\u00a0\u00a0\u00a0\u00a0acute kidney injury\n"
    "Normal text with    four ASCII spaces in the middle.\n"
    "Em-space\u2003\u2003\u2003test and en-space\u2002\u2002\u2002test.\n"
    "Single spaces should stay untouched.\n"
    "Tabs\t\there too."
)

SAMPLE_REFS = [
    {"id": 5, "file": "Paper_A.pdf", "page": 1, "text": PDF_TEXT_WITH_ARTIFACTS},
    {"id": 12, "file": "Paper_A.pdf", "page": 2, "text": "Clean text without artifacts."},
    {"id": 23, "file": "Paper_B.pdf", "page": 1, "text": "Another\t\t\tclean\t\t\ttext."},
    {"id": 30, "file": "Paper_B.pdf", "page": 5, "text": "Uncited source text."},
]

SAMPLE_ANSWER = (
    "Mitochondrial dysfunction interacts with aging hallmarks [5]. "
    "This was further explored in later work [12]. "
    "In contrast, Chen et al. found different results [23]."
)


# ---------------------------------------------------------------------------
# Tests: Whitespace Cleanup
# ---------------------------------------------------------------------------

class TestWhitespaceCleanup:
    """Verifies that TXT export regex handles all whitespace variants from PDF extraction."""

    def test_ascii_tabs_removed(self):
        result = clean_export_text("word\t\tword")
        assert "\t" not in result
        assert result == "word word"

    def test_multiple_ascii_spaces_collapsed(self):
        result = clean_export_text("word    word")
        assert "  " not in result
        assert result == "word word"

    def test_non_breaking_spaces_collapsed(self):
        result = clean_export_text("word\u00a0\u00a0\u00a0word")
        assert "\u00a0" not in result
        assert result == "word word"

    def test_em_spaces_collapsed(self):
        result = clean_export_text("word\u2003\u2003word")
        assert "\u2003" not in result
        assert result == "word word"

    def test_en_spaces_collapsed(self):
        result = clean_export_text("word\u2002\u2002word")
        assert "\u2002" not in result
        assert result == "word word"

    def test_mixed_whitespace_collapsed(self):
        result = clean_export_text("word\t \u00a0\u2003word")
        assert result == "word word"

    def test_single_space_preserved(self):
        result = clean_export_text("word word")
        assert result == "word word"

    def test_single_tab_preserved(self):
        """A single tab is only 1 whitespace char, so it stays (not 2+)."""
        result = clean_export_text("word\tword")
        assert result == "word\tword"

    def test_newlines_preserved(self):
        result = clean_export_text("line one\nline two\n\nline four")
        assert result == "line one\nline two\n\nline four"

    def test_full_pdf_artifact_text(self):
        result = clean_export_text(PDF_TEXT_WITH_ARTIFACTS)
        # No consecutive non-newline whitespace should remain
        assert not re.search(r'[^\S\n]{2,}', result), (
            f"Found consecutive whitespace in cleaned text: {result!r}"
        )
        # Newlines should be preserved
        assert result.count('\n') == PDF_TEXT_WITH_ARTIFACTS.count('\n')


# ---------------------------------------------------------------------------
# Tests: Source Renumbering
# ---------------------------------------------------------------------------

class TestRenumbering:
    """Verifies citation renumbering from internal chunk IDs to sequential display IDs."""

    def test_basic_renumbering(self):
        renumber = build_renumber_map("See [5] and [12] and [23].")
        assert renumber == {5: 1, 12: 2, 23: 3}

    def test_renumber_text_application(self):
        renumber = {5: 1, 12: 2, 23: 3}
        result = renumber_text("Results in [5] confirm [12].", renumber)
        assert result == "Results in [1] confirm [2]."

    def test_uncited_ids_untouched(self):
        renumber = {5: 1}
        result = renumber_text("See [5] and [99].", renumber)
        assert result == "See [1] and [99]."

    def test_empty_answer_no_renumber(self):
        renumber = build_renumber_map("")
        assert renumber == {}

    def test_duplicate_citations_single_entry(self):
        renumber = build_renumber_map("[5] confirms [5] again.")
        assert renumber == {5: 1}

    def test_renumber_preserves_order(self):
        """IDs are renumbered by their numeric value, not by appearance order."""
        renumber = build_renumber_map("[23] appears before [5].")
        assert renumber == {5: 1, 23: 2}


# ---------------------------------------------------------------------------
# Tests: Evidence Filtering
# ---------------------------------------------------------------------------

class TestEvidenceFiltering:
    """Verifies that only cited sources appear in the evidence list."""

    def test_only_cited_refs_shown(self):
        displayed = filter_cited_refs(SAMPLE_REFS, SAMPLE_ANSWER)
        displayed_ids = {r['id'] for r in displayed}
        assert displayed_ids == {5, 12, 23}
        assert 30 not in displayed_ids

    def test_fallback_when_no_citations(self):
        displayed = filter_cited_refs(SAMPLE_REFS, "No citations here.")
        assert len(displayed) == 3  # fallback: first 3

    def test_empty_refs_returns_empty(self):
        displayed = filter_cited_refs([], SAMPLE_ANSWER)
        assert displayed == []


# ---------------------------------------------------------------------------
# Tests: All Sources TXT Export
# ---------------------------------------------------------------------------

class TestAllSourcesTxt:
    """Verifies the 'Download All Sources' TXT bundle structure and content."""

    def test_header_present(self):
        txt = build_all_sources_txt(SAMPLE_REFS[:2])
        assert "DOCUINSIGHT - ALL SOURCES" in txt
        assert "Total sources: 2" in txt

    def test_all_sources_included(self):
        txt = build_all_sources_txt(SAMPLE_REFS)
        for r in SAMPLE_REFS:
            assert f"SOURCE [{r['id']}]" in txt
            assert f"File: {r['file']}" in txt
            assert f"Page: {r['page']}" in txt

    def test_whitespace_cleaned_in_bundle(self):
        txt = build_all_sources_txt(SAMPLE_REFS)
        # The full artifact text should be cleaned in the output
        assert "\t\t" not in txt
        assert "\u00a0\u00a0" not in txt
        assert "\u2003\u2003" not in txt

    def test_newlines_preserved_in_bundle(self):
        refs = [{"id": 1, "file": "test.pdf", "page": 1, "text": "line1\nline2\nline3"}]
        txt = build_all_sources_txt(refs)
        assert "line1\nline2\nline3" in txt


# ---------------------------------------------------------------------------
# Tests: Word Export (ReportGenerator)
# ---------------------------------------------------------------------------

class TestWordExport:
    """Verifies Word report generation without Streamlit."""

    def _make_messages(self, answer: str, refs: list[dict] | None = None) -> list[dict]:
        msgs = [
            {"role": "user", "content": "What are the main findings?"},
            {
                "role": "assistant",
                "content": answer,
                "extended_data": {
                    "final_answer": answer,
                    "references": refs or [],
                },
            },
        ]
        return msgs

    def test_word_generates_bytes(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        msgs = self._make_messages("Simple answer without citations.")
        result = gen.create_word_report("Test Report", msgs)
        assert isinstance(result, io.BytesIO)
        assert result.getvalue()[:2] == b'PK'  # docx is a zip

    def test_word_contains_question_and_answer(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        msgs = self._make_messages("The mitochondria is the powerhouse.")
        result = gen.create_word_report("Test Report", msgs)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "What are the main findings?" in full_text
        assert "mitochondria" in full_text

    def test_word_renumbers_citations(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        answer = "Finding A [5] and finding B [12]."
        refs = [
            {"id": 5, "file": "A.pdf", "page": 1},
            {"id": 12, "file": "A.pdf", "page": 2},
        ]
        msgs = self._make_messages(answer, refs)
        result = gen.create_word_report("Test", msgs)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[1]" in full_text
        assert "[2]" in full_text

    def test_word_includes_source_list(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        answer = "See [5] for details."
        refs = [{"id": 5, "file": "Paper_A.pdf", "page": 3}]
        msgs = self._make_messages(answer, refs)
        result = gen.create_word_report("Test", msgs)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Paper_A.pdf" in full_text
        assert "Page 3" in full_text

    def test_word_empty_messages(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        result = gen.create_word_report("Empty", [])
        assert isinstance(result, io.BytesIO)

    def test_word_skips_system_messages(self):
        from exporter import ReportGenerator
        gen = ReportGenerator()
        msgs = [
            {"role": "system", "content": "You are DocuInsight."},
            {"role": "user", "content": "Hello"},
            {"role": "assistant", "content": "Hi there."},
        ]
        result = gen.create_word_report("Test", msgs)
        doc = Document(result)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "You are DocuInsight" not in full_text
        assert "Hello" in full_text
